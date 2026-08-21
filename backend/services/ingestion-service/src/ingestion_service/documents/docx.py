from __future__ import annotations

import io
from collections.abc import Iterator
from pathlib import Path

from docx import Document
from docx.parts.image import ImagePart
from PIL import Image

from ingestion_service.domain import Asset, BlockType, DocumentUnit, ExtractedBlock


class DOCXReader:
    def supports(self, path: Path, content_type: str) -> bool:
        return path.suffix.lower() == ".docx" or content_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def units(self, path: Path) -> Iterator[DocumentUnit]:
        document = Document(path)
        blocks: list[ExtractedBlock] = []
        section: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name or "").lower() if paragraph.style else ""
            if style.startswith("heading") or style == "title":
                section = [text]
                block_type = BlockType.HEADING
            elif "list" in style:
                block_type = BlockType.LIST
            else:
                block_type = BlockType.PARAGRAPH
            blocks.append(
                ExtractedBlock(block_type=block_type, text=text, section_path=list(section))
            )
        for table in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                text = "\n".join(" | ".join(row) for row in rows)
                blocks.append(
                    ExtractedBlock(
                        block_type=BlockType.TABLE,
                        text=text,
                        section_path=list(section),
                    )
                )
        assets: list[Asset] = []
        seen: set[str] = set()
        for part in document.part.related_parts.values():
            if not isinstance(part, ImagePart):
                continue
            data = part.blob
            with Image.open(io.BytesIO(data)) as image:
                width, height = image.size
                extension = (image.format or "png").lower()
            digest = part.sha1
            if digest in seen:
                continue
            seen.add(digest)
            assets.append(
                Asset(
                    data=data,
                    extension=extension,
                    content_type=part.content_type,
                    width=width,
                    height=height,
                    digest=digest,
                )
            )
        yield DocumentUnit(
            number=1,
            label=path.name,
            blocks=blocks,
            assets=assets,
            metadata={"format": "docx"},
        )
